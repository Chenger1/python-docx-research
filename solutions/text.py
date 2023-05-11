from docx.text.paragraph import Paragraph
from docx.oxml.xmlchemy import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches


# Original Answer: https://stackoverflow.com/questions/48663788/python-docx-insert-a-paragraph-after
def insert_paragraph_after(paragraph, text=None, style=None):
    """Insert a new paragraph after the given paragraph."""
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    if style is not None:
        new_para.style = style
    return new_para


# Original Answer: https://stackoverflow.com/a/70647528/10462999
def table_insert_paragraph_after(table):
    """Return new `Paragraph` object inserted directly after `table`.

    `table` must already be immediately followed by a paragraph. So
    This won't work for a table followed by another table or a table
    at the end of the document.
    """
    p = table._tbl.getnext()
    paragraph = Paragraph(p, table._parent)
    return paragraph.insert_paragraph_before()


# Insert paragraph at specific position
# You have to insert content in reverse order
# Created lots of useless empty paragraphs
def insert_at_position(text, position, base_paragraph):
    paragraph = insert_paragraph_after(base_paragraph, text)
    for i in range(position+1):
        insert_paragraph_after(base_paragraph, "")
    return paragraph


# Delete parapgraph
# Original answer https://github.com/python-openxml/python-docx/issues/33#issuecomment-77661907
# Read comment because solution has some caveats
def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


# DIVIDE Into columns
# Original Answer https://github.com/python-openxml/python-docx/issues/167#issuecomment-772391447
WNS_COLS_NUM = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}num"


def set_number_of_columns(section, cols):
    """ sets number of columns through xpath. """
    section._sectPr.xpath("./w:cols")[0].set(WNS_COLS_NUM, str(cols))


# Original Answer: https://stackoverflow.com/a/59786239
# On some cases you may need to edit functions to apply special styles
def add_bookmark(paragraph, bookmark_text, bookmark_name, run=None):
    # run = paragraph.add_run()
    # tag = run._r  # some versions of editors require run._r instead of xpath
    # tag = paragraph._element.xpath("//w:p")[-1]  # in some cases
    tag = paragraph._element
    start = OxmlElement('w:bookmarkStart')
    start.set(qn('w:id'), '60')  # should be unique for each bookmark
    start.set(qn('w:name'), bookmark_name)
    start.set(qn('w:colFirst'), "0")
    start.set(qn('w:colLast'), "0")
    tag.append(start)

    if run:  # if we have to add bookmark to existing paragraph
        tag.append(run._element)
    else:
        text = OxmlElement('w:r')
        text.text = bookmark_text
        tag.append(text)

    end = OxmlElement('w:bookmarkEnd')
    end.set(qn('w:id'), '60')
    end.set(qn('w:name'), bookmark_name)
    tag.append(end)


# Original Answer: https://stackoverflow.com/a/59786239
def add_link(paragraph, link_to, text, tool_tip=None, run=None):
    # create hyperlink node
    hyperlink = OxmlElement('w:hyperlink')

    # set attribute for link to bookmark
    hyperlink.set(qn('w:anchor'), link_to,)

    if tool_tip is not None:
        # set attribute for link to bookmark
        hyperlink.set(qn('w:tooltip'), tool_tip,)

    if run:   # if we have to add bookmark to existing paragraph
        new_run = run._element
    else:
        new_run = OxmlElement('w:r')
        new_run.text = text
    rPr = OxmlElement('w:rPr')
    new_run.append(rPr)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    # r = paragraph.add_run()
    # r._r.append(hyperlink)


# Set list of items in two columns
# IMPORTANT. This is a hack. For right realisation should make research in open xml specification
def insert_two_columns(list_of_items, prev_paragraph):
    prev_paragraph = prev_paragraph
    for index, item in enumerate(list_of_items, start=1):
        subgroup_paragraph = insert_paragraph_after(prev_paragraph, "")
        subgroup_paragraph.style = "addition_header"
        subgroup_paragraph.add_run(item)
        paragraph_format = subgroup_paragraph.paragraph_format
        paragraph_format.space_after = Pt(2)
        if index % 2 == 0:
            paragraph_format.left_indent = Inches(1)
        else:
            paragraph_format.left_indent = Inches(4)
